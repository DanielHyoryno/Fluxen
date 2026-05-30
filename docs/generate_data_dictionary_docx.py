from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt


OUTPUT_PATH = "docs/data-dictionary.docx"


tables = [
    (
        "Tabel 3.__ Tabel User",
        [
            ("id", "BIGINT", "Primary key tabel user."),
            ("full_name", "VARCHAR", "Nama lengkap pengguna."),
            ("email", "VARCHAR", "Email pengguna, bersifat unik."),
            ("password_hash", "TEXT", "Password pengguna yang telah di-hash."),
            ("role", "VARCHAR", "Role pengguna, default user."),
            ("created_at", "TIMESTAMP", "Waktu data pengguna dibuat."),
            ("updated_at", "TIMESTAMP", "Waktu data pengguna terakhir diperbarui."),
        ],
    ),
    (
        "Tabel 3.__ Tabel Device Category",
        [
            ("id", "BIGINT", "Primary key tabel kategori perangkat."),
            ("user_id", "BIGINT", "Foreign key yang mengacu ke tabel users."),
            ("name", "VARCHAR", "Nama kategori perangkat."),
            ("created_at", "TIMESTAMP", "Waktu data kategori dibuat."),
        ],
    ),
    (
        "Tabel 3.__ Tabel Device",
        [
            ("id", "BIGINT", "Primary key tabel perangkat."),
            ("user_id", "BIGINT", "Foreign key yang mengacu ke tabel users."),
            ("category_id", "BIGINT", "Foreign key yang mengacu ke tabel device_categories."),
            ("device_code", "VARCHAR", "Kode unik perangkat IoT."),
            ("device_name", "VARCHAR", "Nama perangkat IoT."),
            ("api_token_hash", "TEXT", "Token API perangkat yang telah di-hash."),
            ("status", "VARCHAR", "Status koneksi perangkat, default offline."),
            ("install_location", "VARCHAR", "Lokasi pemasangan perangkat."),
            ("firmware_version", "VARCHAR", "Versi firmware perangkat."),
            ("last_seen_at", "TIMESTAMP", "Waktu terakhir perangkat aktif atau mengirim data."),
            ("created_at", "TIMESTAMP", "Waktu data perangkat dibuat."),
            ("updated_at", "TIMESTAMP", "Waktu data perangkat terakhir diperbarui."),
        ],
    ),
    (
        "Tabel 3.__ Tabel Measurement",
        [
            ("id", "BIGINT", "Primary key tabel pengukuran."),
            ("device_id", "BIGINT", "Foreign key yang mengacu ke tabel devices."),
            ("measured_at", "TIMESTAMP", "Waktu pengukuran dilakukan."),
            ("flow_rate_lpm", "NUMERIC", "Debit air dalam satuan liter per menit."),
            ("volume_delta_l", "NUMERIC", "Volume air yang terukur pada interval tertentu dalam liter."),
            ("cumulative_volume_l", "NUMERIC", "Total volume kumulatif yang tercatat oleh perangkat."),
            ("pulse_count", "INTEGER", "Jumlah pulsa sensor flow meter."),
            ("battery_voltage", "NUMERIC", "Tegangan baterai perangkat."),
            ("rssi_dbm", "INTEGER", "Kekuatan sinyal perangkat dalam satuan dBm."),
            ("created_at", "TIMESTAMP", "Waktu data pengukuran disimpan."),
        ],
    ),
    (
        "Tabel 3.__ Tabel Device Threshold",
        [
            ("id", "BIGINT", "Primary key tabel batas perangkat."),
            ("device_id", "BIGINT", "Foreign key unik yang mengacu ke tabel devices."),
            ("leak_flow_min_lpm", "NUMERIC", "Batas minimum debit air untuk indikasi kebocoran."),
            ("leak_duration_sec", "INTEGER", "Durasi minimum aliran air untuk indikasi kebocoran."),
            ("quiet_start_time", "TIME", "Waktu mulai periode pemantauan senyap."),
            ("quiet_end_time", "TIME", "Waktu akhir periode pemantauan senyap."),
            ("daily_usage_limit_l", "NUMERIC", "Batas penggunaan air harian dalam liter."),
            ("monthly_usage_limit_l", "NUMERIC", "Batas penggunaan air bulanan dalam liter."),
            ("created_at", "TIMESTAMP", "Waktu data batas perangkat dibuat."),
            ("updated_at", "TIMESTAMP", "Waktu data batas perangkat terakhir diperbarui."),
        ],
    ),
    (
        "Tabel 3.__ Tabel Alert",
        [
            ("id", "BIGINT", "Primary key tabel peringatan."),
            ("device_id", "BIGINT", "Foreign key yang mengacu ke tabel devices."),
            ("alert_type", "VARCHAR", "Jenis peringatan yang dihasilkan sistem."),
            ("severity", "VARCHAR", "Tingkat keparahan peringatan, default medium."),
            ("title", "VARCHAR", "Judul peringatan."),
            ("message", "TEXT", "Isi pesan peringatan."),
            ("triggered_at", "TIMESTAMP", "Waktu peringatan dipicu."),
            ("resolved_at", "TIMESTAMP", "Waktu peringatan diselesaikan."),
            ("status", "VARCHAR", "Status peringatan, default active."),
            ("meta", "JSONB", "Metadata tambahan terkait peringatan."),
            ("created_at", "TIMESTAMP", "Waktu data peringatan dibuat."),
        ],
    ),
    (
        "Tabel 3.__ Tabel Bill Settings",
        [
            ("id", "BIGINT", "Primary key tabel pengaturan tagihan."),
            ("user_id", "BIGINT", "Foreign key unik yang mengacu ke tabel users."),
            ("price_per_liter", "NUMERIC", "Harga air per liter yang digunakan untuk estimasi tagihan."),
            ("currency", "VARCHAR", "Mata uang yang digunakan, default IDR."),
            ("created_at", "TIMESTAMP", "Waktu data pengaturan tagihan dibuat."),
            ("updated_at", "TIMESTAMP", "Waktu data pengaturan tagihan terakhir diperbarui."),
        ],
    ),
]


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


document = Document()
section = document.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

heading = document.add_paragraph("3.3.4.2 Data Dictionary")
heading.style = document.styles["Heading 2"]

for title, rows in tables:
    paragraph = document.add_paragraph(title)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in paragraph.runs:
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)

    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True

    headers = ["Atribut", "Tipe", "Keterangan"]
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)

    for attribute, data_type, description in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], attribute)
        set_cell_text(cells[1], data_type)
        set_cell_text(cells[2], description)

    document.add_paragraph("")

document.save(OUTPUT_PATH)
print(OUTPUT_PATH)
